from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


TODAY = date(2026, 3, 23)
OUTPUT_DIR = Path("output/doc/legal")

COMPANY = {
    "app_name": "Pure Grief and Therapeutic ART",
    "operator": "De Troostboom",
    "website": "https://pure-therapeutic-art-therapy.com",
    "email": "info@pure-therapeutic-art-therapy.com",
    "kvk": "17251011",
    "address_line_1": "[IN TE VULLEN: vestigingsadres De Troostboom]",
    "address_line_2": "[IN TE VULLEN: postcode en vestigingsplaats]",
    "country": "Nederland",
    "vat": "[IN TE VULLEN: btw-nummer indien van toepassing]",
    "phone": "[OPTIONEEL IN TE VULLEN: telefoonnummer]",
    "privacy_email": "info@pure-therapeutic-art-therapy.com",
}


def iso_date() -> str:
    return TODAY.strftime("%d-%m-%Y")


def add_paragraph(doc: Document, text: str, *, italic: bool = False) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.italic = italic


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.5)

    for style_name, size in [("Title", 20), ("Heading 1", 14), ("Heading 2", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
        style.font.size = Pt(size)


def write_document(definition: dict[str, Any]) -> Path:
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(definition["title"])

    subtitle = definition.get("subtitle")
    if subtitle:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(subtitle)
        run.italic = True

    meta_lines = definition.get("meta_lines", [])
    for line in meta_lines:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(line)

    doc.add_paragraph()

    for section in definition["sections"]:
        doc.add_heading(section["heading"], level=1)

        for paragraph in section.get("paragraphs", []):
            add_paragraph(doc, paragraph)

        for subheading, subparagraphs in section.get("subsections", []):
            doc.add_heading(subheading, level=2)
            for paragraph in subparagraphs:
                add_paragraph(doc, paragraph)

        bullets = section.get("bullets")
        if bullets:
            add_bullets(doc, bullets)

        if section.get("snippet"):
            doc.add_heading("Klaar-om-te-gebruiken tekst", level=2)
            add_paragraph(doc, section["snippet"], italic=True)

        if section.get("note"):
            add_paragraph(doc, f"Let op: {section['note']}", italic=True)

    path = OUTPUT_DIR / definition["filename"]
    doc.save(path)
    return path


def definitions() -> list[dict[str, Any]]:
    year = TODAY.year
    return [
        {
            "filename": "00_Overzicht_juridisch_pakket_Pure_Grief_and_Therapeutic_ART.docx",
            "title": "Overzicht juridisch pakket",
            "subtitle": f"{COMPANY['app_name']} - opleverpakket juridische documenten",
            "meta_lines": [
                f"Exploitant: {COMPANY['operator']}",
                f"Datum: {iso_date()}",
                "Versie: 1.0",
            ],
            "sections": [
                {
                    "heading": "Doel van dit pakket",
                    "paragraphs": [
                        "Dit pakket bevat de juridische documenten en implementatieteksten voor de app en website van Pure Grief and Therapeutic ART. De documenten zijn zo opgezet dat ze direct bruikbaar zijn als basis voor publicatie in de app, op de website en in ondersteunende schermen zoals registratie, de hulpverlenerskaart en werkvormen.",
                        "De documenten zijn inhoudelijk aangescherpt op basis van jullie bestaande concepten en aangevuld met onderdelen die nog ontbraken, zoals een duidelijkere privacystructuur, expliciete digitale-contentclausules, abonnementsregels, licentiebeperkingen en praktische teksten voor overige pagina's in de app.",
                    ],
                },
                {
                    "heading": "Documenten in dit pakket",
                    "bullets": [
                        "01 Algemene voorwaarden en condities - hoofdovereenkomst voor app en bijbehorende digitale diensten.",
                        "02 Privacyverklaring AVG - verwerking van persoonsgegevens, rechten van gebruikers en gegevensdeling.",
                        "03 Disclaimer - korte juridische grens van de app en de werkvormen.",
                        "04 Impressum - formele bedrijfs- en publicatiegegevens.",
                        "05 Copyright en licentievoorwaarden - auteursrecht, gebruikslicentie en verboden hergebruik.",
                        "06 Platformdisclaimer hulpverlenerskaart - rol van het platform en afbakening ten opzichte van de therapeutenkaart.",
                        "07 Community & Professional Guidelines - regels voor hulpverleners op het platform.",
                        "08 Registratie- en akkoordteksten - checkboxteksten, korte Terms of Use en toestemmingsblokken.",
                        "09 Verklaring hulpverlener kaartpublicatie - aparte verklaring voor therapeuten die zichtbaar willen zijn op de kaart.",
                        "10 Teksten voor werkvormen, e-books en appstores - standaardteksten voor andere plekken in de app en store listings.",
                    ],
                },
                {
                    "heading": "Plaatsing in app en website",
                    "bullets": [
                        "Veiligheid & privacy in de accountomgeving: Disclaimer, Algemene voorwaarden en condities, Privacyverklaring AVG, Impressum en Copyright.",
                        "Registratieflow: checkboxtekst en korte Terms of Use uit document 08.",
                        "Hulpverlenerskaart: platformdisclaimer uit document 06 en kaartverklaring uit document 09.",
                        "Werkvormen en e-books: copyright- en licentieteksten uit document 10.",
                        "Website footer: links naar Algemene voorwaarden, Privacyverklaring, Disclaimer, Impressum en Copyright.",
                    ],
                },
                {
                    "heading": "Nog handmatig te bevestigen gegevens",
                    "bullets": [
                        f"Vestigingsadres: {COMPANY['address_line_1']}, {COMPANY['address_line_2']}.",
                        f"BTW-nummer: {COMPANY['vat']}.",
                        f"Telefoonnummer of extra contactkanaal: {COMPANY['phone']}.",
                        "Definitieve lijst van verwerkers en subverwerkers in de privacyverklaring, voor zover je die expliciet wilt benoemen.",
                        "Eventuele specifieke retour- of klachtenprocedure voor aankopen via de externe website van De Troostboom.",
                    ],
                    "note": "Laat de definitieve documenten voor publicatie nog toetsen door een Nederlandse jurist of advocaat, met name op consumentenrecht, abonnementen, AVG en eventuele internationale doorgifte van persoonsgegevens.",
                },
            ],
        },
        {
            "filename": "01_Algemene_voorwaarden_en_condities_Pure_Grief_and_Therapeutic_ART.docx",
            "title": "Algemene voorwaarden en condities",
            "subtitle": COMPANY["app_name"],
            "meta_lines": [
                f"Exploitant: {COMPANY['operator']}",
                f"KvK: {COMPANY['kvk']}",
                f"Website: {COMPANY['website']}",
                f"Laatst bijgewerkt: {iso_date()}",
            ],
            "sections": [
                {
                    "heading": "1. Identiteit van de aanbieder",
                    "paragraphs": [
                        f"De applicatie {COMPANY['app_name']} en de bijbehorende website worden ontwikkeld en geëxploiteerd door {COMPANY['operator']}, gevestigd te {COMPANY['address_line_1']}, {COMPANY['address_line_2']}, {COMPANY['country']}, ingeschreven bij de Kamer van Koophandel onder nummer {COMPANY['kvk']}.",
                        f"Voor vragen over deze voorwaarden kan contact worden opgenomen via {COMPANY['email']}.",
                    ],
                },
                {
                    "heading": "2. Toepasselijkheid",
                    "paragraphs": [
                        "Deze algemene voorwaarden zijn van toepassing op ieder gebruik van de app, de website, de therapeutenkaart, de digitale content, downloads, e-books, accounts, credits, abonnementen en overige aan de app gekoppelde diensten van de aanbieder.",
                        "Door een account aan te maken, content te openen, een aankoop te doen of een abonnement af te sluiten, verklaart de gebruiker zich met deze voorwaarden akkoord.",
                    ],
                },
                {
                    "heading": "3. Definities",
                    "bullets": [
                        "App: de digitale omgeving Pure Grief and Therapeutic ART.",
                        "Gebruiker: iedere natuurlijke persoon of rechtspersoon die gebruikmaakt van de app of website.",
                        "Hulpverlener: een gebruiker die als therapeut, begeleider, coach, docent of andere professional diensten aanbiedt of zichtbaar is op de therapeutenkaart.",
                        "Content: alle werkvormen, hoofdstukken, e-books, teksten, illustraties, afbeeldingen, audio, downloads, formats en overige materialen in of via de app.",
                        "Credits: digitale tegoeden waarmee toegang tot bepaalde content kan worden verkregen.",
                        "Abonnement: een periodieke betaalde toegangsvorm voor diensten of content in de app.",
                    ],
                },
                {
                    "heading": "4. Doel en aard van de app",
                    "paragraphs": [
                        "De app biedt een digitale omgeving met creatieve, therapeutische en educatieve werkvormen, hoofdstukken, e-books en ondersteunende functies rond verlies, rouw, trauma, emotionele verwerking, cognitie, gedrag, zintuiglijke beleving, natuurbeleving en persoonlijke ontwikkeling.",
                        "De inhoud van de app is bedoeld voor reflectie, educatie en ondersteuning. De aanbieder verleent via de app geen medische, psychologische of therapeutische behandeling.",
                    ],
                },
                {
                    "heading": "5. Gebruik en professionele verantwoordelijkheid",
                    "paragraphs": [
                        "Gebruikers blijven te allen tijde zelf verantwoordelijk voor hun keuzes, interpretaties en het gebruik van content uit de app. Voor hulpverleners geldt dat zij zelf verantwoordelijk blijven voor diagnostiek, behandelkeuzes, methodische afwegingen, cliëntveiligheid en beroepsethisch handelen.",
                        "Indien een gebruiker de inhoud inzet binnen een professionele praktijk, geschiedt dat volledig onder de eigen professionele verantwoordelijkheid van die gebruiker.",
                    ],
                },
                {
                    "heading": "6. Account, registratie en beveiliging",
                    "bullets": [
                        "Een gebruiker is verantwoordelijk voor de juistheid van de bij registratie verstrekte gegevens.",
                        "Een gebruiker moet inloggegevens vertrouwelijk behandelen en mag het account niet zonder toestemming door anderen laten gebruiken.",
                        "De aanbieder mag het account tijdelijk blokkeren of beëindigen bij misbruik, fraude, inbreuk op intellectuele eigendomsrechten, schending van deze voorwaarden of gedragingen die schade kunnen veroorzaken aan de app, andere gebruikers of derden.",
                    ],
                },
                {
                    "heading": "7. Digitale content, aankopen, credits en levering",
                    "paragraphs": [
                        "Toegang tot bepaalde content, hoofdstukken, e-books of functies kan afhankelijk zijn van aankoop, ontgrendeling via credits, of een actief abonnement. De aanbieder bepaalt welke onderdelen gratis toegankelijk zijn en welke onderdelen achter een betaalmuur vallen.",
                        "Na een geslaagde betaling of geldige ontgrendeling wordt de betreffende digitale content gekoppeld aan het account van de gebruiker. De aanbieder spant zich in om die toegang direct of zo spoedig mogelijk beschikbaar te maken.",
                    ],
                    "bullets": [
                        "Aangekochte credits zijn in beginsel niet restitueerbaar.",
                        "Digitale content wordt geleverd door toegang in de app of via het account van de gebruiker.",
                        "Websitebestellingen die gekoppeld zijn aan het app-account kunnen in het accountoverzicht zichtbaar worden zodra de technische koppeling dat ondersteunt.",
                    ],
                },
                {
                    "heading": "8. Abonnementen, verlenging en opzegging",
                    "paragraphs": [
                        "Een abonnement geeft gedurende de overeengekomen looptijd toegang tot de functies of content die op het moment van aankoop bij dat abonnement horen.",
                        "Indien een abonnement stilzwijgend wordt verlengd, gebeurt dat uitsluitend onder de voorwaarden die op grond van het toepasselijke consumentenrecht zijn toegestaan. Voor consumenten geldt dat zij na een stilzwijgende verlenging het abonnement moeten kunnen beëindigen met een opzegtermijn van maximaal één maand, voor zover de wet dat vereist.",
                        "De gebruiker kan een abonnement beheren, verlenging uitzetten of opzeggen via het daarvoor aangewezen accountscherm of via de route die de aanbieder daarvoor beschikbaar stelt.",
                    ],
                },
                {
                    "heading": "9. Herroepingsrecht en digitale inhoud",
                    "paragraphs": [
                        "Voor zover een gebruiker kwalificeert als consument, kan voor aankopen op afstand in beginsel een wettelijk herroepingsrecht gelden. Voor digitale inhoud die niet op een materiële drager wordt geleverd en die direct toegankelijk wordt gemaakt, kan dat herroepingsrecht vervallen zodra de levering begint, mits de consument daar vooraf uitdrukkelijk mee instemt en erkent dat het herroepingsrecht daarmee vervalt, voor zover de wet dat toestaat.",
                        "De aanbieder dient deze instemming duidelijk in het aankoopproces vast te leggen wanneer sprake is van directe levering van digitale content.",
                    ],
                },
                {
                    "heading": "10. Therapeutenkaart en platformfunctie",
                    "paragraphs": [
                        "De app kan een kaart, directory of overzicht bevatten waarin hulpverleners zichtbaar zijn voor potentiële cliënten. De aanbieder fungeert daarbij uitsluitend als technisch platform en is geen bemiddelaar, zorgaanbieder of contractpartij bij de relatie tussen cliënt en hulpverlener.",
                        "Vermelding op de kaart betekent geen aanbeveling, certificering, kwaliteitsgarantie of inhoudelijke toetsing van de betreffende hulpverlener.",
                    ],
                },
                {
                    "heading": "11. Intellectueel eigendom en licentie",
                    "paragraphs": [
                        f"Alle rechten op de app en de content berusten bij {COMPANY['operator']} of bij de betreffende rechthebbenden. De gebruiker ontvangt uitsluitend een beperkte, niet-exclusieve, niet-overdraagbare en herroepbare licentie om de content te gebruiken binnen het doel waarvoor deze in de app beschikbaar is gesteld.",
                    ],
                    "bullets": [
                        "Het is niet toegestaan om content te kopiëren, te publiceren, door te verkopen, systematisch te verzamelen, te delen buiten de toegestane context of zonder toestemming te verwerken in trainingen, cursussen, methodieken of eigen commerciële producten.",
                        "Voor e-books en andere beschermde digitale inhoud geldt dat deze uitsluitend binnen de daartoe aangewezen app- of readeromgeving mogen worden gebruikt, tenzij uitdrukkelijk anders is aangegeven.",
                    ],
                },
                {
                    "heading": "12. Verboden gebruik en handhaving",
                    "bullets": [
                        "Misleidende, onrechtmatige, discriminerende of schadelijke handelingen via de app zijn verboden.",
                        "Het omzeilen van technische beveiliging, kopieerbeperkingen, toegangsrechten of betalingsstromen is verboden.",
                        "De aanbieder mag bij een redelijk vermoeden van misbruik content ontoegankelijk maken, accounts beperken, licenties intrekken en nadere maatregelen nemen.",
                    ],
                },
                {
                    "heading": "13. Aansprakelijkheid",
                    "paragraphs": [
                        "Gebruik van de app en de content gebeurt op eigen risico. De aanbieder is, voor zover wettelijk toegestaan, niet aansprakelijk voor directe of indirecte schade, gevolgschade, emotionele of psychologische schade, lichamelijke gevolgen, financiële schade, verlies van inkomsten, verlies van reputatie of claims van derden die voortvloeien uit het gebruik van de app, de content of het contact tussen gebruikers en hulpverleners.",
                        "De uitsluiting van aansprakelijkheid geldt niet voor zover schade het gevolg is van opzet of grove nalatigheid van de aanbieder en niet verder dan wettelijk is toegestaan.",
                        "Indien aansprakelijkheid toch wordt vastgesteld, is deze beperkt tot het bedrag dat de gebruiker in de twaalf maanden voorafgaand aan de gebeurtenis aan de aanbieder heeft betaald voor het gebruik van de app.",
                    ],
                },
                {
                    "heading": "14. Vrijwaring",
                    "paragraphs": [
                        "De gebruiker vrijwaart de aanbieder voor aanspraken van derden, waaronder cliënten, mede-gebruikers en zakelijke relaties, voor zover die aanspraken voortvloeien uit het gebruik van de app, de toepassing van werkvormen, de professionele activiteiten van de gebruiker of schending van deze voorwaarden.",
                    ],
                },
                {
                    "heading": "15. Privacy en persoonsgegevens",
                    "paragraphs": [
                        "Persoonsgegevens worden verwerkt conform de Privacyverklaring AVG van de aanbieder. De gebruiker dient deze privacyverklaring gelijktijdig met deze voorwaarden te raadplegen.",
                    ],
                },
                {
                    "heading": "16. Wijzigingen, klachten en toepasselijk recht",
                    "paragraphs": [
                        "De aanbieder mag deze voorwaarden aanpassen wanneer wetgeving, diensten of de werking van de app daarom vragen. De meest recente versie wordt via de app en/of website gepubliceerd.",
                        f"Klachten over de app of het gebruik ervan kunnen worden gemeld via {COMPANY['email']}. De aanbieder streeft ernaar klachten binnen een redelijke termijn inhoudelijk te beantwoorden.",
                        "Op deze voorwaarden is Nederlands recht van toepassing. Geschillen worden, behoudens dwingendrechtelijke uitzonderingen, voorgelegd aan de bevoegde rechter in Nederland.",
                    ],
                },
            ],
        },
        {
            "filename": "02_Privacyverklaring_AVG_Pure_Grief_and_Therapeutic_ART.docx",
            "title": "Privacyverklaring AVG",
            "subtitle": COMPANY["app_name"],
            "meta_lines": [
                f"Verwerkingsverantwoordelijke: {COMPANY['operator']}",
                f"Laatst bijgewerkt: {iso_date()}",
            ],
            "sections": [
                {
                    "heading": "1. Wie verantwoordelijk is voor de verwerking",
                    "paragraphs": [
                        f"{COMPANY['operator']} is verantwoordelijk voor de verwerking van persoonsgegevens in de app en op de website van {COMPANY['app_name']}.",
                        f"Contactgegevens: {COMPANY['address_line_1']}, {COMPANY['address_line_2']}, {COMPANY['country']}, e-mail: {COMPANY['privacy_email']}, KvK: {COMPANY['kvk']}.",
                    ],
                },
                {
                    "heading": "2. Welke persoonsgegevens worden verwerkt",
                    "bullets": [
                        "Accountgegevens, zoals naam, e-mailadres, taalvoorkeur, accountstatus en inloggegevens.",
                        "Profielgegevens, zoals profielfoto, praktijkinformatie, beroepsgegevens, openbare contactgegevens en informatie voor de therapeutenkaart.",
                        "Aankoop- en abonnementsgegevens, zoals bestelhistorie, credits, ontgrendelde content, factuur- en orderreferenties en abonnementsstatus.",
                        "Gebruiksgegevens, zoals bezochte pagina's, voortgang in content, voorkeuren, app-interacties en technische loggegevens.",
                        "Communicatiegegevens, zoals berichten aan support, e-mailverkeer, testmails of bevestigingen.",
                        "Locatie- en zichtbaarheidgegevens voor therapeuten die ervoor kiezen op de kaart of in een directory zichtbaar te zijn.",
                        "Vrije tekst die een gebruiker zelf invult, zoals notities of profielinformatie. De app is niet bedoeld voor het opslaan van medische dossiers of andere bijzondere persoonsgegevens, tenzij de gebruiker daar zelf bewust en rechtmatig voor kiest.",
                    ],
                },
                {
                    "heading": "3. Voor welke doelen en op welke grondslag",
                    "subsections": [
                        (
                            "Account en authenticatie",
                            [
                                "Doel: het aanmaken, beveiligen en beheren van accounts, inloggen, sessiebeheer en toegangscontrole.",
                                "Grondslag: uitvoering van de overeenkomst en gerechtvaardigd belang bij beveiliging van het platform.",
                            ],
                        ),
                        (
                            "Levering van digitale content en abonnementen",
                            [
                                "Doel: verwerking van aankopen, toegang tot content, abonnementen, credits, ontgrendelingen en accountgebonden producten.",
                                "Grondslag: uitvoering van de overeenkomst en, waar nodig, wettelijke administratieverplichtingen.",
                            ],
                        ),
                        (
                            "Therapeutenkaart en openbare profielen",
                            [
                                "Doel: zichtbaarheid van hulpverleners voor potentiële cliënten en vindbaarheid binnen de app of website.",
                                "Grondslag: toestemming en uitvoering van de overeenkomst met de betreffende hulpverlener.",
                            ],
                        ),
                        (
                            "Communicatie en ondersteuning",
                            [
                                "Doel: beantwoorden van vragen, versturen van accountmails, serviceberichten en belangrijke updates.",
                                "Grondslag: uitvoering van de overeenkomst, gerechtvaardigd belang en soms toestemming.",
                            ],
                        ),
                        (
                            "Verbetering, beveiliging en analyse",
                            [
                                "Doel: analyseren van gebruik, verbeteren van functionaliteit, voorkomen van misbruik en beveiligen van de app.",
                                "Grondslag: gerechtvaardigd belang en, voor zover wettelijk vereist, toestemming voor niet strikt noodzakelijke tracking of analytics.",
                            ],
                        ),
                    ],
                },
                {
                    "heading": "4. Van wie de gegevens afkomstig zijn",
                    "paragraphs": [
                        "Persoonsgegevens worden in beginsel rechtstreeks van de gebruiker ontvangen, bijvoorbeeld bij registratie, profielinvulling, gebruik van de app of bij een aankoop.",
                        "Daarnaast kunnen gegevens afkomstig zijn van externe systemen die door of namens de gebruiker worden gebruikt, zoals de gekoppelde website, betaalproviders of orderkoppelingen, voor zover die integraties onderdeel zijn van de dienstverlening.",
                    ],
                },
                {
                    "heading": "5. Met wie gegevens worden gedeeld",
                    "paragraphs": [
                        "De aanbieder verkoopt geen persoonsgegevens. Gegevens kunnen wel worden gedeeld met verwerkers of dienstverleners die nodig zijn voor het functioneren van de app.",
                    ],
                    "bullets": [
                        "Authenticatie-, database- en opslagproviders, waaronder Supabase voor account-, database- en mediaverwerking.",
                        "E-mail- en communicatieproviders, waaronder Google/Gmail of vergelijkbare e-maildiensten voor service- en notificatiemails.",
                        "Hosting- en serverbeheerdiensten voor het beschikbaar houden van de app en website.",
                        "Betaal- en orderverwerkers voor het afhandelen van betalingen, abonnementen en koppelingen met de externe webshop of WooCommerce-omgeving.",
                        "Analysediensten of loggingtools voor prestaties, foutanalyse en gebruiksstatistieken, voor zover deze zijn ingeschakeld.",
                    ],
                },
                {
                    "heading": "6. Internationale doorgifte",
                    "paragraphs": [
                        "Indien persoonsgegevens worden verwerkt door dienstverleners buiten de Europese Economische Ruimte, zorgt de aanbieder voor passende waarborgen, zoals een adequaatheidsbesluit van de Europese Commissie of standaardcontractbepalingen, voor zover wettelijk vereist.",
                    ],
                },
                {
                    "heading": "7. Bewaartermijnen",
                    "bullets": [
                        "Accountgegevens: gedurende de looptijd van het account en daarna zolang nodig is voor afhandeling, beveiliging of wettelijke verplichtingen.",
                        "Aankoop- en administratieve gegevens: zolang nodig voor uitvoering van de overeenkomst en de wettelijke fiscale bewaarplicht.",
                        "Therapeutenprofielen en kaartgegevens: zolang het profiel actief is en daarna nog kort voor afwikkeling, back-up en bewijsdoeleinden.",
                        "Support- en contactverzoeken: zolang nodig om de vraag af te handelen en redelijke tijd daarna voor dossieropbouw.",
                        "Voortgangs-, voorkeur- en gebruiksgegevens: zolang het account actief is of totdat de gebruiker deze gegevens verwijdert, voor zover de functionaliteit dat toelaat.",
                    ],
                },
                {
                    "heading": "8. Beveiliging",
                    "paragraphs": [
                        "De aanbieder neemt passende technische en organisatorische maatregelen om persoonsgegevens te beveiligen tegen verlies, ongeautoriseerde toegang, misbruik of onrechtmatige verwerking. Denk aan toegangsbeperking, versleutelde verbindingen, logging, scheiding van rollen en back-upmaatregelen.",
                    ],
                },
                {
                    "heading": "9. Rechten van betrokkenen",
                    "paragraphs": [
                        "Gebruikers hebben, voor zover de AVG dat toelaat, recht op inzage, rectificatie, verwijdering, beperking van de verwerking, dataportabiliteit en bezwaar. Verzoeken kunnen worden ingediend via het contactadres van de aanbieder.",
                        "De aanbieder reageert in beginsel binnen één maand op een verzoek. Indien dat nodig is, kan deze termijn binnen de wettelijke kaders worden verlengd.",
                    ],
                },
                {
                    "heading": "10. Klacht indienen",
                    "paragraphs": [
                        f"Indien een gebruiker vindt dat de aanbieder persoonsgegevens niet correct verwerkt, kan eerst contact worden opgenomen via {COMPANY['privacy_email']}. Daarnaast heeft een gebruiker het recht een klacht in te dienen bij de Autoriteit Persoonsgegevens.",
                        "Autoriteit Persoonsgegevens, Postbus 93374, 2509 AJ Den Haag, Nederland.",
                    ],
                },
                {
                    "heading": "11. Cookies, sessies en analytics",
                    "paragraphs": [
                        "De webapp en website kunnen functionele cookies of vergelijkbare technieken gebruiken voor inloggen, sessiebeheer, beveiliging, taalvoorkeuren, splash-/voorkeursinstellingen en het correct laten werken van de omgeving.",
                        "Daarnaast kunnen analytische of vergelijkbare meetinstrumenten worden ingezet om het gebruik van de app te begrijpen en te verbeteren. Voor niet strikt noodzakelijke cookies of tracking vraagt de aanbieder, voor zover wettelijk vereist, vooraf toestemming.",
                    ],
                },
                {
                    "heading": "12. Minderjarigen",
                    "paragraphs": [
                        "De app richt zich primair op volwassenen en professionals. Voor zover de app door minderjarigen wordt gebruikt, dient dit te gebeuren met toestemming van een ouder, voogd of andere bevoegde vertegenwoordiger, voor zover de wet dat vereist.",
                    ],
                },
                {
                    "heading": "13. Wijzigingen",
                    "paragraphs": [
                        "De aanbieder kan deze privacyverklaring aanpassen wanneer wetgeving, techniek of dienstverlening verandert. De meest actuele versie wordt via de app en/of website beschikbaar gemaakt.",
                    ],
                },
            ],
        },
        {
            "filename": "03_Disclaimer_Pure_Grief_and_Therapeutic_ART.docx",
            "title": "Disclaimer",
            "subtitle": COMPANY["app_name"],
            "meta_lines": [f"Laatst bijgewerkt: {iso_date()}"],
            "sections": [
                {
                    "heading": "1. Doel van de app",
                    "paragraphs": [
                        f"{COMPANY['app_name']} biedt creatieve, educatieve en ondersteunende content rondom verlies, rouw, trauma, emotionele verwerking, persoonlijke ontwikkeling en aanverwante thema's.",
                    ],
                },
                {
                    "heading": "2. Geen vervanging van behandeling of spoedzorg",
                    "paragraphs": [
                        "De app, de werkvormen, e-books, hoofdstukken en overige materialen vormen geen medische, psychologische, psychiatrische of therapeutische behandeling en vervangen geen professionele diagnostiek, crisisinterventie of spoedhulp.",
                        "Bij acute psychische nood, suïcidaliteit, medische klachten of andere urgente situaties moet direct contact worden opgenomen met een arts, crisisdienst, 112 of een andere passende hulpverlener.",
                    ],
                },
                {
                    "heading": "3. Eigen verantwoordelijkheid van de gebruiker",
                    "paragraphs": [
                        "Gebruik van de inhoud van de app gebeurt volledig onder eigen verantwoordelijkheid van de gebruiker. Hulpverleners, therapeuten en begeleiders blijven zelf verantwoordelijk voor hun beroepsmatig handelen, de geschiktheid van een werkvorm, de timing van toepassing en de veiligheid van cliënten of deelnemers.",
                    ],
                },
                {
                    "heading": "4. Geen garantie op resultaat",
                    "paragraphs": [
                        "De aanbieder garandeert niet dat een werkvorm, methode of oefening geschikt, effectief of veilig is voor iedere persoon, context of hulpvraag. De inhoud van de app kan verwijzen naar verschillende methodieken, stromingen of modellen, maar dat betekent geen garantie op een bepaald therapeutisch of persoonlijk resultaat.",
                    ],
                },
                {
                    "heading": "5. Aansprakelijkheid",
                    "paragraphs": [
                        f"{COMPANY['operator']} is, voor zover wettelijk toegestaan, niet aansprakelijk voor schade of nadelige gevolgen die voortvloeien uit de interpretatie of toepassing van de inhoud van de app, noch voor handelen van derden die via de app of de therapeutenkaart zichtbaar worden gemaakt.",
                    ],
                },
            ],
        },
        {
            "filename": "04_Impressum_Pure_Grief_and_Therapeutic_ART.docx",
            "title": "Impressum",
            "subtitle": COMPANY["app_name"],
            "meta_lines": [f"Laatst bijgewerkt: {iso_date()}"],
            "sections": [
                {
                    "heading": "Juridische gegevens",
                    "bullets": [
                        f"Exploitant / verwerkingsverantwoordelijke: {COMPANY['operator']}",
                        f"Handelsnaam / productnaam: {COMPANY['app_name']}",
                        f"Vestigingsadres: {COMPANY['address_line_1']}",
                        f"Postcode en plaats: {COMPANY['address_line_2']}",
                        f"Land: {COMPANY['country']}",
                        f"E-mailadres: {COMPANY['email']}",
                        f"Telefoon: {COMPANY['phone']}",
                        f"KvK-nummer: {COMPANY['kvk']}",
                        f"BTW-nummer: {COMPANY['vat']}",
                        f"Website: {COMPANY['website']}",
                    ],
                },
                {
                    "heading": "Verantwoordelijke voor inhoud en publicatie",
                    "paragraphs": [
                        f"Tenzij anders vermeld, is {COMPANY['operator']} verantwoordelijk voor de inhoud en publicatie van de app, de website en de via de app aangeboden juridische en informatieve teksten.",
                    ],
                },
                {
                    "heading": "Contact voor juridische kennisgevingen",
                    "paragraphs": [
                        f"Voor juridische kennisgevingen, klachten, auteursrechtmeldingen, privacyverzoeken of andere formele correspondentie kan contact worden opgenomen via {COMPANY['email']}. Indien de aanbieder een postadres voor formele kennisgevingen wenst te gebruiken, dient dit adres in de bovenstaande gegevens definitief te worden ingevuld.",
                    ],
                },
            ],
        },
        {
            "filename": "05_Copyright_en_licentievoorwaarden_Pure_Grief_and_Therapeutic_ART.docx",
            "title": "Copyright en licentievoorwaarden",
            "subtitle": COMPANY["app_name"],
            "meta_lines": [f"Laatst bijgewerkt: {iso_date()}"],
            "sections": [
                {
                    "heading": "1. Rechten op content",
                    "paragraphs": [
                        f"Alle werkvormen, hoofdstukken, e-books, downloads, teksten, visuele materialen, opmaak, databases, methodische structuren en overige content in of via de app zijn beschermd door auteursrecht en andere intellectuele eigendomsrechten. Deze rechten berusten bij {COMPANY['operator']} of bij de betreffende rechthebbenden.",
                    ],
                },
                {
                    "heading": "2. Gebruikslicentie",
                    "paragraphs": [
                        "De gebruiker ontvangt uitsluitend een beperkte, niet-exclusieve, niet-overdraagbare en herroepbare licentie om de content te gebruiken binnen de context waarvoor die is aangeschaft of beschikbaar is gesteld.",
                    ],
                    "bullets": [
                        "Professionals mogen werkvormen gebruiken binnen hun eigen praktijk of beroepsuitoefening, voor zover de app of aankoop dat toestaat.",
                        "De licentie geeft geen eigendomsrecht en geen recht om content door te verkopen, te sublicentiëren of te publiceren.",
                        "E-books en beveiligde digitale content mogen uitsluitend in de door de aanbieder aangewezen omgeving worden geraadpleegd, tenzij schriftelijk anders is toegestaan.",
                    ],
                },
                {
                    "heading": "3. Verboden handelingen",
                    "bullets": [
                        "Het kopiëren, herpubliceren, uploaden, doorverkopen of verspreiden van content buiten de app of buiten de toegestane professionele context.",
                        "Het opnemen van werkvormen of materialen in eigen trainingen, cursussen, boeken, online modules of commerciële producten zonder voorafgaande schriftelijke toestemming.",
                        "Het verwijderen of omzeilen van technische beschermingsmaatregelen, watermerken, toegangsbeperkingen of readerbeveiligingen.",
                        "Het systematisch verzamelen van content uit de app, bijvoorbeeld via scraping, automatisering of bulkdownloads.",
                    ],
                },
                {
                    "heading": "4. Beveiligde content en e-books",
                    "paragraphs": [
                        "Voor e-books en beveiligde digitale content mag de aanbieder technische maatregelen inzetten om ongeoorloofd kopiëren, printen, delen of opslaan te beperken. Het bestaan van dergelijke maatregelen geeft de gebruiker geen recht om deze te omzeilen.",
                    ],
                },
                {
                    "heading": "5. Notice-and-takedown / inbreukmelding",
                    "paragraphs": [
                        f"Wie meent dat content in of rond de app inbreuk maakt op zijn of haar auteursrecht of andere intellectuele eigendomsrechten, kan daarvan melding maken via {COMPANY['email']}. De melding moet voldoende concreet zijn, zodat de aanbieder de klacht kan beoordelen en zo nodig passende maatregelen kan nemen.",
                    ],
                },
                {
                    "heading": "6. Standaard copyrightregel voor publicatie",
                    "snippet": f"© {COMPANY['operator']} - {COMPANY['app_name']}. Alle rechten voorbehouden. Gebruik binnen de eigen praktijk is uitsluitend toegestaan voor de geregistreerde gebruiker. Kopiëren, delen, publiceren of doorverkopen zonder schriftelijke toestemming is verboden.",
                },
            ],
        },
        {
            "filename": "06_Platformdisclaimer_hulpverlenerskaart_Pure_Grief_and_Therapeutic_ART.docx",
            "title": "Platformdisclaimer hulpverlenerskaart",
            "subtitle": COMPANY["app_name"],
            "meta_lines": [f"Laatst bijgewerkt: {iso_date()}"],
            "sections": [
                {
                    "heading": "1. Rol van het platform",
                    "paragraphs": [
                        f"De therapeutenkaart, hulpverlenerskaart of directory binnen {COMPANY['app_name']} is uitsluitend bedoeld als vindbaarheids- en zichtbaarheidsfunctie. {COMPANY['operator']} faciliteert de technische publicatie van profielen, maar is geen bemiddelaar, zorgaanbieder of contractspartij bij contacten tussen cliënten en hulpverleners.",
                    ],
                },
                {
                    "heading": "2. Geen kwaliteitsgarantie",
                    "paragraphs": [
                        "Vermelding op de kaart betekent geen aanbeveling, certificering, verificatie of kwaliteitsgarantie. De aanbieder controleert niet actief alle diploma's, registraties, behandelmethoden, prijzen of resultaten van vermelde hulpverleners.",
                    ],
                },
                {
                    "heading": "3. Zelfstandige verantwoordelijkheid van hulpverleners",
                    "paragraphs": [
                        "Iedere hulpverlener op de kaart werkt zelfstandig en is volledig verantwoordelijk voor zijn of haar eigen beroepsuitoefening, methodiek, cliëntcontact, intake, tarieven, overeenkomst, privacy, dossiervorming en naleving van lokale wet- en regelgeving.",
                    ],
                },
                {
                    "heading": "4. Geen aansprakelijkheid voor behandelrelaties",
                    "paragraphs": [
                        f"{COMPANY['operator']} is niet aansprakelijk voor schade, klachten, geschillen, behandeluitkomsten, financiële afspraken of andere gevolgen die voortvloeien uit het contact tussen cliënten en hulpverleners die via de kaart of directory zichtbaar zijn.",
                    ],
                },
                {
                    "heading": "5. Tekst boven de kaart",
                    "snippet": f"De hulpverleners op deze kaart zijn zelfstandige professionals en werken onafhankelijk van {COMPANY['operator']}. {COMPANY['operator']} biedt uitsluitend een platform voor zichtbaarheid en vindbaarheid en geeft geen kwaliteitsgarantie of behandeladvies.",
                },
                {
                    "heading": "6. Korte waarschuwing onder of naast de kaart",
                    "snippet": "Vermelding op deze kaart betekent geen aanbeveling, certificering of kwaliteitsgarantie door de aanbieder van de app.",
                },
            ],
        },
        {
            "filename": "07_Community_en_Professional_Guidelines_Pure_Grief_and_Therapeutic_ART.docx",
            "title": "Community & Professional Guidelines",
            "subtitle": f"Voor hulpverleners en professionele gebruikers van {COMPANY['app_name']}",
            "meta_lines": [f"Laatst bijgewerkt: {iso_date()}"],
            "sections": [
                {
                    "heading": "1. Professioneel handelen",
                    "paragraphs": [
                        "Professionele gebruikers en hulpverleners die deelnemen aan de app of de therapeutenkaart handelen volgens de voor hun vakgebied geldende professionele en ethische normen en volgens de wet- en regelgeving die in hun land of werkgebied van toepassing is.",
                    ],
                },
                {
                    "heading": "2. Correcte en actuele profielinformatie",
                    "bullets": [
                        "Profielinformatie moet waarheidsgetrouw, actueel en controleerbaar zijn.",
                        "Misleidende claims over opleiding, registratie, behandelresultaten, methoden of specialisaties zijn niet toegestaan.",
                        "Contactgegevens, werkgebied en zichtbaarheidinstellingen moeten zo actueel mogelijk worden gehouden.",
                    ],
                },
                {
                    "heading": "3. Respectvolle communicatie",
                    "bullets": [
                        "Gebruikers communiceren respectvol met cliënten, andere hulpverleners en de aanbieder.",
                        "Discriminatie, intimidatie, misleiding, agressie of ander grensoverschrijdend gedrag is niet toegestaan.",
                        "Het gebruiken van de app voor spam, acquisitie buiten de context van de app of ongepaste promotie is niet toegestaan.",
                    ],
                },
                {
                    "heading": "4. Privacy en vertrouwelijkheid",
                    "paragraphs": [
                        "Hulpverleners zijn zelf verantwoordelijk voor vertrouwelijke omgang met cliëntinformatie en mogen via publieke profielen of open velden geen gegevens publiceren die zij niet rechtmatig mogen delen. De app is geen vervanging van een professioneel cliëntdossier of een beveiligde behandelomgeving.",
                    ],
                },
                {
                    "heading": "5. Grenzen van gebruik van content",
                    "paragraphs": [
                        "Werkvormen en materialen uit de app mogen uitsluitend worden gebruikt binnen de grenzen van de verleende licentie en op een manier die past bij de deskundigheid en beroepsverantwoordelijkheid van de gebruiker. Onjuist of onveilig gebruik, herpublicatie en commercieel hergebruik zonder toestemming zijn niet toegestaan.",
                    ],
                },
                {
                    "heading": "6. Meldingen, handhaving en verwijdering",
                    "paragraphs": [
                        f"{COMPANY['operator']} behoudt zich het recht voor om profielen, accounts of content te weigeren, te beperken of te verwijderen wanneer deze richtlijnen worden overtreden, wanneer misleidende informatie wordt geplaatst, wanneer klachten of veiligheidsrisico's ontstaan of wanneer de integriteit van het platform in het geding komt.",
                    ],
                },
            ],
        },
        {
            "filename": "08_Registratie_en_akkoordteksten_Pure_Grief_and_Therapeutic_ART.docx",
            "title": "Registratie- en akkoordteksten",
            "subtitle": COMPANY["app_name"],
            "meta_lines": [f"Laatst bijgewerkt: {iso_date()}"],
            "sections": [
                {
                    "heading": "1. Korte Terms of Use bij accountregistratie",
                    "snippet": f"Door een account aan te maken in {COMPANY['app_name']} verklaar ik dat ik de Algemene voorwaarden, Privacyverklaring en Disclaimer heb gelezen en accepteer. Ik begrijp dat de app educatieve en ondersteunende content bevat en geen medische of therapeutische behandeling vervangt. Ik blijf zelf verantwoordelijk voor de manier waarop ik de inhoud gebruik.",
                },
                {
                    "heading": "2. Checkboxtekst bij registratie",
                    "snippet": "Ik ga akkoord met de Algemene voorwaarden, Privacyverklaring en Disclaimer van De Troostboom.",
                },
                {
                    "heading": "3. Extra checkbox bij directe levering van digitale content",
                    "snippet": "Ik stem ermee in dat de levering van digitale content direct begint en erken dat mijn herroepingsrecht vervalt zodra de levering is gestart, voor zover de wet dat toestaat.",
                },
                {
                    "heading": "4. Extra tekst voor professionele gebruikers",
                    "snippet": "Ik verklaar dat ik de inhoud van de app op een professionele, zorgvuldige en verantwoordelijke manier gebruik en volledig verantwoordelijk ben voor mijn eigen handelen en eventuele toepassing in de praktijk.",
                },
            ],
        },
        {
            "filename": "09_Verklaring_hulpverlener_kaartpublicatie_Pure_Grief_and_Therapeutic_ART.docx",
            "title": "Verklaring hulpverlener voor kaartpublicatie",
            "subtitle": COMPANY["app_name"],
            "meta_lines": [f"Laatst bijgewerkt: {iso_date()}"],
            "sections": [
                {
                    "heading": "Verklaring",
                    "paragraphs": [
                        "Door mijn profiel zichtbaar te maken op de hulpverlenerskaart of in een openbare directory van de app verklaar ik dat:",
                    ],
                    "bullets": [
                        "ik zelfstandig werk en niet optreed namens De Troostboom;",
                        "ik beschik over relevante opleiding, ervaring of kwalificaties voor mijn werkzaamheden;",
                        "ik volledig verantwoordelijk ben voor mijn eigen beroepsuitoefening, behandelkeuzes, ethiek en communicatie met cliënten;",
                        "de informatie in mijn profiel correct, actueel en niet misleidend is;",
                        "ik zelf verantwoordelijk ben voor naleving van de wet- en regelgeving die op mijn beroep en werkgebied van toepassing is;",
                        "ik De Troostboom vrijwaar voor claims of geschillen die voortvloeien uit mijn werkzaamheden of mijn contact met cliënten.",
                    ],
                },
                {
                    "heading": "Checkboxtekst",
                    "snippet": "Ik verklaar dat ik zelfstandig en professioneel handel, dat mijn profielinformatie juist is en dat ik volledig verantwoordelijk ben voor mijn beroepsuitoefening en cliëntcontacten.",
                },
            ],
        },
        {
            "filename": "10_Teksten_voor_werkvormen_ebooks_en_appstores_Pure_Grief_and_Therapeutic_ART.docx",
            "title": "Teksten voor werkvormen, e-books en appstores",
            "subtitle": COMPANY["app_name"],
            "meta_lines": [f"Laatst bijgewerkt: {iso_date()}"],
            "sections": [
                {
                    "heading": "1. Standaardtekst onder iedere werkvorm",
                    "snippet": f"Deze werkvorm is onderdeel van {COMPANY['app_name']} en is bedoeld voor professionele toepassing of zorgvuldig persoonlijk gebruik binnen de context van de app. © {COMPANY['operator']} - gebruik binnen de eigen praktijk of het eigen account is toegestaan voor de geregistreerde gebruiker. Verdere verspreiding, publicatie of doorverkoop is niet toegestaan.",
                },
                {
                    "heading": "2. Korte disclaimer onder werkvormen",
                    "snippet": "Deze inhoud is informatief en ondersteunend van aard en vervangt geen professionele diagnostiek, behandeling of medische zorg. Gebruik gebeurt onder eigen verantwoordelijkheid.",
                },
                {
                    "heading": "3. E-book readertekst",
                    "snippet": "Dit e-book is uitsluitend beschikbaar binnen de beveiligde app-omgeving. Opslaan, printen, kopiëren, delen of commercieel hergebruik buiten de app is niet toegestaan zonder schriftelijke toestemming van De Troostboom.",
                },
                {
                    "heading": "4. Tekst op product- of aankooppagina voor digitale content",
                    "snippet": "Na aankoop wordt deze digitale content direct aan uw account gekoppeld. Voor zover wettelijk toegestaan en nadat u daar expliciet mee hebt ingestemd, vervalt het herroepingsrecht zodra de levering van de digitale inhoud is gestart.",
                },
                {
                    "heading": "5. Korte app store / Play Store beschrijving",
                    "snippet": f"{COMPANY['app_name']} biedt creatieve, educatieve en ondersteunende werkvormen, hoofdstukken en e-books voor verlies, rouw en therapeutische reflectie. De app is geen vervanging van medische of therapeutische behandeling en werkt aanvullend op professioneel en zorgvuldig gebruik.",
                },
                {
                    "heading": "6. Tekst boven de therapeutenkaart",
                    "snippet": f"De hulpverleners op deze kaart werken onafhankelijk en zijn volledig verantwoordelijk voor hun eigen professionele handelen. {COMPANY['operator']} biedt uitsluitend een platform voor zichtbaarheid.",
                },
            ],
        },
    ]


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    paths = [write_document(definition) for definition in definitions()]
    print("Generated documents:")
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
